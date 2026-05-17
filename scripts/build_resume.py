#!/usr/bin/env python
"""生成谢琨鹏-西北工业大学的技术简历"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


def set_cell_border(cell, **kwargs):
    """设置表格单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), '0')
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=None):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing


def add_run(paragraph, text, bold=False, size=None, color=None, font_name=None, font_name_east=None):
    """添加格式化文本块"""
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    if font_name_east:
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), font_name_east)
        rPr.insert(0, rFonts)
    return run


def build_resume():
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # ── 样式设置 ──
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)
    style.font.name = 'Arial'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 姓名标题 ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=4)
    add_run(p, '谢琨鹏', bold=True, size=Pt(22), font_name='Arial', font_name_east='黑体')

    # ── 联系方式 ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=2)
    contact = ('电话：xxxx-xxxx-xxxx  |  邮箱：xxxx@xxxx.com  |  '
               'GitHub：https://github.com/xxxx  |  求职意向：C++ 后端开发工程师')
    add_run(p, contact, size=Pt(9), color=RGBColor(0x55, 0x55, 0x55))

    # ── 分隔线 ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=2)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '1F4E79')
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ══════════════════════════════════════════
    # 教育背景
    # ══════════════════════════════════════════
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    add_run(p, '教育背景', bold=True, size=Pt(13), color=RGBColor(0x1F, 0x4E, 0x79), font_name='Arial', font_name_east='黑体')

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, '西北工业大学', bold=True, size=Pt(10.5), font_name='Arial', font_name_east='宋体')
    add_run(p, '　　　　　　　　　　　　　　　　　　　　　　xxxx年xx月 - xxxx年xx月', size=Pt(9), color=RGBColor(0x66, 0x66, 0x66))

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, '计算机科学与技术 / 软件工程  |  本科/硕士  |  GPA：x.x / 4.0', size=Pt(9), color=RGBColor(0x55, 0x55, 0x55))

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, '主修课程：数据结构与算法、操作系统、计算机网络、数据库系统、计算机组成原理、C++程序设计等', size=Pt(9), color=RGBColor(0x55, 0x55, 0x55))

    # ══════════════════════════════════════════
    # 技术栈
    # ══════════════════════════════════════════
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    add_run(p, '技术栈', bold=True, size=Pt(13), color=RGBColor(0x1F, 0x4E, 0x79), font_name='Arial', font_name_east='黑体')

    skills_table = doc.add_table(rows=6, cols=2)
    skills_table.autofit = True

    skill_data = [
        ('编程语言', 'C/C++（主力）、Python、SQL'),
        ('标准库 & STL', '精通 STL 容器与算法内部原理，能独立实现红黑树、哈希表、内存池等核心数据结构'),
        ('框架 & 工具', 'Qt 5.x GUI 开发、OpenCV 图像处理、MySQL 数据库、CMake 构建系统'),
        ('底层知识', '内存管理（内存池、自由链表）、迭代器模式、类型萃取（type traits）、异常安全保证'),
        ('开发工具', 'Visual Studio、Git、GCC/G++、CMake'),
        ('其他', 'Linux/Windows 跨平台开发、多线程编程基础、设计模式'),
    ]

    for i, (label, content) in enumerate(skill_data):
        row = skills_table.rows[i]

        # 设置行高
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcHeight = OxmlElement('w:tcHeight')
            tcHeight.set(qn('w:val'), '360')
            tcHeight.set(qn('w:hRule'), 'atLeast')
            tcPr.append(tcHeight)

        # 第一列
        row.cells[0].width = Cm(2.5)
        p = row.cells[0].paragraphs[0]
        add_run(p, label, bold=True, size=Pt(9.5), font_name='Arial', font_name_east='宋体')

        # 第二列
        row.cells[1].width = Cm(13.5)
        p = row.cells[1].paragraphs[0]
        add_run(p, content, size=Pt(9.5), font_name='Arial', font_name_east='宋体')

    # 去掉表格边框
    for row in skills_table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                            top={'val': 'nil'}, bottom={'val': 'nil'},
                            left={'val': 'nil'}, right={'val': 'nil'})

    # ══════════════════════════════════════════
    # 项目经历
    # ══════════════════════════════════════════
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10, after=4)
    add_run(p, '项目经历', bold=True, size=Pt(13), color=RGBColor(0x1F, 0x4E, 0x79), font_name='Arial', font_name_east='黑体')

    # ── 项目 1：MyTinySTL ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=0)
    add_run(p, 'MyTinySTL — C++11 微型 STL 标准库实现', bold=True, size=Pt(11), font_name='Arial', font_name_east='宋体')
    add_run(p, '　　C++11 / 数据结构 / 内存管理', size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, '项目描述：', bold=True, size=Pt(9.5))
    add_run(p, '基于 C++11 标准，从零实现了一套完整的微型 STL 库（mystl），涵盖标准库中的核心容器、算法与内存管理组件，'
             '所有容器与算法均支持迭代器操作，并编写了完整的单元测试套件。', size=Pt(9.5))

    items_mystl = [
        '内存管理子系统：实现了二级空间配置器 alloc，采用内存池 + 自由链表（56 个链表覆盖 8~4096 字节），'
        '通过 M_round_up 对齐与 M_chunk_alloc 按需分配策略优化小内存分配效率；实现了 construct / destroy 的完美转发与类型萃取优化',
        '序列式容器：实现了 vector（动态扩容，支持强异常安全保证）、deque（双端队列，基于中控器 + 分段连续空间）、'
        'list（双向循环链表）、stack / queue（基于 deque 的适配器模式）',
        '关联式容器：从底层数据结构出发，完全实现了红黑树 rb_tree（含插入再平衡、迭代器自增自减）和哈希表 hashtable（开链法冲突解决），'
        '并在此基础上构建了 map、set、unordered_map、unordered_set',
        '算法库：实现了 sort（内省排序）、heap 操作系列、set 交并差算法、numeric 数值算法等，均支持自定义比较器与迭代器范围',
        '类型系统：实现了 type_traits（类型萃取）、iterator（五类迭代器 + traits）、functional（函数对象）等基础设施',
        '跨平台构建：基于 CMake 组织项目，兼容 Linux（g++ 5.4+）、Windows（msvc 14.0+）、macOS（clang 3.5+）三平台',
    ]

    for item in items_mystl:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)
        pf = p.paragraph_format
        pf.left_indent = Cm(0.5)
        add_run(p, '▸ ', size=Pt(9.5), color=RGBColor(0x1F, 0x4E, 0x79))
        add_run(p, item, size=Pt(9.5))

    # ── 项目 2：Telemedicine ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=0)
    add_run(p, 'Telemedicine — 远程医疗影像辅助诊断系统', bold=True, size=Pt(11), font_name='Arial', font_name_east='宋体')
    add_run(p, '　　Qt 5 / OpenCV / MySQL / C++', size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, '项目描述：', bold=True, size=Pt(9.5))
    add_run(p, '开发了一款基于 Qt 的桌面端远程医疗影像辅助诊断系统，集成患者信息管理、CT 影像处理与病灶自动识别功能，'
             '旨在辅助医生进行子宫肌瘤的初步筛查诊断。', size=Pt(9.5))

    items_tele = [
        '数据库管理层：基于 MySQL 8.0 设计 patient 数据库，包含 user_profile 表（患者基本信息 + 病例 + 照片）、'
        'basic_inf 和 details_inf 两个视图；利用 QSqlTableModel 实现患者数据的界面化增删改查，'
        '支持按姓名联动查询详细信息与照片；照片以 Base64 编码存储于 BLOB 字段，实现数据与文件系统解耦',
        'CT 影像处理流程：使用 OpenCV 实现完整的 CT 图像处理管线 — 图像读取 → BGR→RGB→Gray 色彩空间转换 → '
        'GaussianBlur 高斯滤波降噪 → HoughCircles 霍夫圆检测（梯度法）→ 在原始图像上绘制检测圆标记可疑病灶区域，'
        '支持 JPG/PNG/BMP 多格式导入与结果保存',
        'GUI 交互设计：基于 Qt Designer 布局主界面（843×691），包含 QTabWidget 双标签页（基本信息 + 病例详情）、'
        'QTableView 患者列表、实时时钟 QTimer、进度条 QProgressBar 反馈处理进度、LCD 数字日期显示，'
        '支持患者信息联动选择与自动年龄计算',
        '容错机制：启动时自动检测 MySQL 连接状态，若连接失败则尝试启动 mysqld 服务进程后重连，'
        '连接失败时弹出中文错误提示并安全退出',
    ]

    for item in items_tele:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)
        pf = p.paragraph_format
        pf.left_indent = Cm(0.5)
        add_run(p, '▸ ', size=Pt(9.5), color=RGBColor(0x1F, 0x4E, 0x79))
        add_run(p, item, size=Pt(9.5))

    # ══════════════════════════════════════════
    # 个人总结
    # ══════════════════════════════════════════
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    add_run(p, '个人总结', bold=True, size=Pt(13), color=RGBColor(0x1F, 0x4E, 0x79), font_name='Arial', font_name_east='黑体')

    summary_items = [
        '具备扎实的 C++ 编程功底，熟悉 C++11/14 标准，理解 RAII、智能指针、移动语义、模板元编程等核心机制',
        '深入理解 STL 底层实现原理（内存分配器、红黑树、哈希表、迭代器体系），具备独立实现核心数据结构的能力',
        '熟悉 Qt 桌面应用开发与 OpenCV 图像处理，具备将算法落地为可交互 GUI 工具的全流程开发经验',
        '掌握 MySQL 数据库设计与开发，熟悉 SQL 查询、视图、索引设计及 C++ 数据库接口编程',
        '具备良好的工程素养：CMake 构建系统、跨平台兼容性设计、异常安全保证、代码可测试性',
        '有较强的自学能力与问题解决能力，对 C++ 底层机制和系统编程有浓厚兴趣',
    ]

    for item in summary_items:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)
        pf = p.paragraph_format
        pf.left_indent = Cm(0.5)
        add_run(p, '• ', size=Pt(9.5), color=RGBColor(0x1F, 0x4E, 0x79))
        add_run(p, item, size=Pt(9.5))

    # ── 保存 ──
    output_path = r'C:\Graduation project\LowDiff\谢琨鹏-西北工业大学.docx'
    doc.save(output_path)
    print(f'简历已保存到：{output_path}')


if __name__ == '__main__':
    build_resume()
